import re
from docx import Document
from markdown2 import markdown
from config import OUTPUT_DIR

class Generator:
    def __init__(self, markdown_text):
        self.markdown_text = markdown_text
        self.document = Document()

    def markdown_to_html(self):
        """Converts Markdown text to HTML."""
        return markdown(self.markdown_text)

    def add_html_to_docx(self, html_text):
        """Parse HTML and add it to the DOCX document, handling headers and paragraphs."""
        # A simple regex to capture header tags and paragraphs
        parts = re.findall(r'<(h[1-6]|p)>(.*?)</\1>', html_text, re.DOTALL)

        for tag, content in parts:
            content = re.sub(r'<[^>]+>', '', content)  # Remove any inner HTML tags
            if tag == 'p':
                self.document.add_paragraph(content)
            elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                style = 'Heading ' + tag[1]  # Example: 'Heading 1' for 'h1'
                self.document.add_paragraph(content, style=style)

    def generate_docx(self, filename):
        """Generates the DOCX file."""
        html_content = self.markdown_to_html()
        self.add_html_to_docx(html_content)
        self.document.save(f"{OUTPUT_DIR}/{filename}")
        return f"Document '{filename}' has been created successfully."

# Example usage:
if __name__ == "__main__":
    markdown_text = """
# Example Heading
This is a simple paragraph in **Markdown**!
"""
    generator = Generator(markdown_text)
    print(generator.generate_docx("output.docx"))
